import json
from typing import Dict, Any
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import io

class FileHandler:
    @staticmethod
    def save_json(data: Dict[str, Any], filename: str) -> None:
        """Save data to a JSON file."""
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f'Error saving JSON to {filename}:', e)
    
    @staticmethod
    def load_json(filename: str) -> Dict[str, Any]:
        """Load data from a JSON file."""
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f'Error loading JSON from {filename}:', e)
    
    @staticmethod
    def create_word_document(content: Dict[str, Any], filename: str) -> None:
        """Create a Word document with the analysis results."""
        try:
            doc = docx.Document()
            
            # Add title
            title = doc.add_heading('Resume Analysis Results', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add sections
            for section, data in content.items():
                # Add section heading
                doc.add_heading(section.replace('_', ' ').title(), level=1)
                
                if isinstance(data, dict):
                    for key, value in data.items():
                        # Add subsection
                        doc.add_heading(key.replace('_', ' ').title(), level=2)
                        
                        if isinstance(value, list):
                            for item in value:
                                doc.add_paragraph(item, style='List Bullet')
                        else:
                            doc.add_paragraph(str(value))
                elif isinstance(data, list):
                    for item in data:
                        doc.add_paragraph(item, style='List Bullet')
                else:
                    doc.add_paragraph(str(data))
                
                # Add spacing between sections
                doc.add_paragraph()
            
            # Save document
            doc.save(filename)
        except Exception as e:
            print(f'Error creating Word document to {filename}:', e)
    
    @staticmethod
    def extract_text_from_pdf(pdf_file: bytes) -> str:
        """Extract text from a PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_file))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text()
            
            return text
        except Exception as e:
            print('Error extracting text from PDF:', e)
    
    @staticmethod
    def extract_text_from_docx(docx_file: bytes) -> str:
        """Extract text from a Word document."""
        try:
            doc = docx.Document(io.BytesIO(docx_file))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
        except Exception as e:
            print('Error extracting text from DOCX:', e)
    
    @staticmethod
    def create_cover_letter_doc(cover_letter: str, filename: str) -> None:
        """Create a Word document with the cover letter."""
        try:
            doc = docx.Document()
            
            # Add title
            title = doc.add_heading('Cover Letter', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add cover letter content
            paragraphs = cover_letter.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save document
            doc.save(filename)
        except Exception as e:
            print(f'Error creating cover letter document to {filename}:', e)
    
    @staticmethod
    def create_tailored_resume_doc(resume_sections: Dict[str, Any], filename: str) -> None:
        """Create a Word document with the tailored resume."""
        try:
            doc = docx.Document()
            
            # Add title
            title = doc.add_heading('Tailored Resume', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add sections
            for section, content in resume_sections.items():
                # Add section heading
                doc.add_heading(section.replace('_', ' ').title(), level=1)
                
                if isinstance(content, list):
                    for item in content:
                        p = doc.add_paragraph(item, style='List Bullet')
                else:
                    p = doc.add_paragraph(str(content))
                
                # Add spacing between sections
                doc.add_paragraph()
            
            # Save document
            doc.save(filename)
        except Exception as e:
            print(f'Error creating tailored resume document to {filename}:', e) 